import io
import tkinter as tk
from typing import List
from pathlib import Path
from tkinter import filedialog, messagebox
from tkinter import ttk
from docx import Document
from docx.shared import Inches
from PIL import Image
import fitz


class BarcodeConverter:
    def __init__(self, root: tk.Tk) -> None:
        self.root = root
        self.input_file_paths: List[str] = []
        self.output_file_path: str = ""
        self.create_ui()

    def create_ui(self) -> None:
        input_file_label = ttk.Label(self.root, text="Input files:")
        input_file_label.grid(column=0, row=0, padx=5, pady=5, sticky=tk.W)

        self.input_file_listbox = tk.Listbox(self.root, height=10, width=60)
        self.input_file_listbox.grid(column=1, row=0, padx=5, pady=5, sticky=tk.W)

        add_file_button = ttk.Button(self.root, text="Add files", command=self.browse_input_files)
        add_file_button.grid(column=2, row=0, padx=5, pady=5, sticky=tk.W)

        output_file_label = ttk.Label(self.root, text="Output file:")
        output_file_label.grid(column=0, row=1, padx=5, pady=5, sticky=tk.W)

        self.output_file_entry = ttk.Entry(self.root, width=60)
        self.output_file_entry.grid(column=1, row=1, padx=5, pady=5, sticky=tk.W)

        output_file_button = ttk.Button(self.root, text="Browse...", command=self.browse_output_file)
        output_file_button.grid(column=2, row=1, padx=5, pady=5, sticky=tk.W)

        output_file_label = ttk.Label(self.root, text="Progress: ")
        output_file_label.grid(column=0, row=3, padx=5, pady=5, sticky=tk.W)

        self.input_progress_bar = ttk.Progressbar(self.root, orient="horizontal", length=367, mode="determinate")
        self.input_progress_bar.grid(column=1, row=3, padx=5, pady=5, sticky=tk.W)

        convert_button = ttk.Button(self.root, text="Convert", command=self.convert_pages)
        convert_button.grid(column=2, row=3, padx=5, pady=5, sticky=tk.W)

    def browse_input_files(self) -> None:
        """Allow the user to select multiple input PDF files"""
        file_types = (("PDF files", "*.pdf"), ("All files", "*.*"))
        file_paths = filedialog.askopenfilenames(parent=self.root, title="Choose input files", filetypes=file_types)
        if file_paths:
            self.input_file_paths = list(file_paths)
            self.input_file_listbox.delete(0, tk.END)
            for file_path in self.input_file_paths:
                self.input_file_listbox.insert(tk.END, file_path)

    def browse_output_file(self) -> None:
        """Allow the user to select the output DOCX file"""
        file_types = (("Word files", "*.docx"), ("All files", "*.*"))
        file_path = filedialog.asksaveasfilename(parent=self.root, title="Save output file", filetypes=file_types,
                                                 defaultextension='.docx')
        if file_path:
            self.output_file_path = file_path
            self.output_file_entry.delete(0, tk.END)
            self.output_file_entry.insert(0, self.output_file_path)

    def convert_pages(self) -> None:
        input_file_paths = self.input_file_paths
        output_file_path = self.output_file_path

        # Check if there are any input files selected
        if not input_file_paths:
            tk.messagebox.showerror("Error", "No input files selected.")
            return

        # Check if the output file path has been set
        if not output_file_path:
            tk.messagebox.showerror("Error", "Output file path not set.")
            return

        try:
            # Create a new output document and table
            output_doc = Document()
            table = output_doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'

            # Get the first row of the table and the number of columns
            row_cells = table.rows[0].cells
            num_cols = len(row_cells)

            curr_row = 0
            curr_col = 0

            # Check if there are enough rows in the table
            if len(table.rows) <= len(input_file_paths):
                for i in range(len(table.rows), len(input_file_paths) + 1):
                    table.add_row()

            # Configure progress bar
            self.input_progress_bar["maximum"] = len(input_file_paths)
            self.input_progress_bar["value"] = 0

            for input_file_path in input_file_paths:
                # Check if the input file exists
                if not Path(input_file_path).is_file():
                    tk.messagebox.showerror("Error", f"Input file not found: {input_file_path}")
                    continue

                # Open the PDF file using PyMuPDF (fitz)
                try:
                    doc = fitz.open(input_file_path)
                except Exception as e:
                    tk.messagebox.showerror("Error", f"Error opening input file {input_file_path}: {e}")
                    continue

                # Process each page in the file
                for page in doc:
                    try:
                        zoom = 300 / 72  # calculate zoom level to get 300 DPI resolution
                        pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom))
                        img = Image.frombytes("RGB", tuple([pix.width, pix.height]), pix.samples)

                        img_bytes = io.BytesIO()
                        img.save(img_bytes, format='PNG')

                        # Add the image to the current cell
                        if curr_row < len(table.rows):
                            p = row_cells[curr_col].add_paragraph()
                            p.add_run().add_picture(img_bytes, width=Inches(1.7), height=Inches(1.2))

                            curr_col += 1

                        else:
                            table.add_row()
                            row_cells = table.rows[curr_row].cells
                            num_cols = len(row_cells)
                            p = row_cells[curr_col].add_paragraph()
                            p.add_run().add_picture(img_bytes, width=Inches(2), height=Inches(1.5))

                            curr_col += 1
                            curr_row += 1

                        # If the maximum number of columns has been reached, move to the next row
                        if curr_col == num_cols:
                            curr_col = 0
                            curr_row += 1

                            # Check if there are enough rows in the table
                            if len(table.rows) <= curr_row:
                                table.add_row()

                            row_cells = table.rows[curr_row].cells

                        # Update the progress bar
                        self.input_progress_bar["value"] += 1
                        self.root.update()

                    except Exception as e:
                        tk.messagebox.showerror("Something went wrong")

            # Save the output file
            output_doc.save(output_file_path)

            messagebox.showinfo("Success", "Conversion complete.")

        except Exception as e:
            messagebox.showerror("Error", f"An error occurred: {e}")


if __name__ == '__main__':
    # create the Tkinter root object
    GUI = tk.Tk()
    GUI.title("Barcode Converter")
    icon_path = "barcodes_icon.ico"
    GUI.iconbitmap(icon_path)

    # create an instance of BarcodeConverter, passing the root as an argument
    converter = BarcodeConverter(GUI)

    try:
        # start the Tkinter event loop
        GUI.mainloop()
    except Exception as e:
        print(f"Error during mainloop: {e}")
